from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Titre
title = doc.add_heading('QUESTIONNAIRE — ÉTUDE DE MARCHÉ IMMOBILIER', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('À destination des agences immobilières de Côte d\'Ivoire')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].italic = True
subtitle.runs[0].font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph('Cette enquête est menée dans le cadre du développement d\'une solution de gestion immobilière adaptée aux réalités ivoiriennes. Vos réponses nous permettront de construire un outil qui répond réellement à vos besoins.').italic = True
doc.add_paragraph('Durée estimée : 10 minutes | Confidentialité assurée')
doc.add_paragraph()

# Section A
doc.add_heading('SECTION A — PROFIL DE L\'AGENCE', level=1)

data_a = [
    ['A1', 'Nom de l\'agence :', '___________________________________________'],
    ['A2', 'Ville / Commune :', '___________________________________________'],
    ['A3', 'Années d\'existence :', '□ < 1 an  □ 1-3 ans  □ 3-5 ans  □ 5-10 ans  □ > 10 ans'],
    ['A4', 'Nombre d\'employés :', '□ 1-2  □ 3-5  □ 6-10  □ > 10'],
    ['A5', 'Volume de biens en location :', '□ < 10  □ 10-30  □ 30-100  □ > 100'],
    ['A6', 'Logiciel actuellement utilisé :', '□ Aucun (Excel/Papier)  □ Logiciel générique  □ Logiciel immobilier : _____________'],
    ['A7', 'Budget mensuel outils numériques :', '□ < 50 000 FCFA  □ 50 000-150 000  □ 150 000-300 000  □ > 300 000'],
]

for item in data_a:
    p = doc.add_paragraph()
    p.add_run(f"{item[0]}. ").bold = True
    p.add_run(item[1])
    p.add_run(f"  {item[2]}")

# Section B
doc.add_page_break()
doc.add_heading('SECTION B — FRUSTRATIONS & PROCESS', level=1)

doc.add_paragraph().add_run('B1. Classez vos 3 plus grosses frustrations (1 = plus frustrant) :').bold = True
frustrations = [
    'Gestion des loyers et quittances',
    'Suivi des baux et échéances',
    'Organisation des visites',
    'Gestion des travaux de rénovation',
    'Gestion syndic (copropriété)',
    'Comptabilité et taxes (TLPPU)',
    'Suivi des ventes et commissions',
    'Manque de visibilité sur la rentabilité'
]
for f in frustrations:
    doc.add_paragraph(f'□ {f}    Rang : _____', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph().add_run('B2. Comment gérez-vous les quittances et le suivi des impayés aujourd\'hui ?').bold = True
doc.add_paragraph('_______________________________________________________________________________________')
doc.add_paragraph('_______________________________________________________________________________________')

doc.add_paragraph()
doc.add_paragraph().add_run('B3. Comment suivez-vous les travaux de rénovation avant mise en location/vente ?').bold = True
doc.add_paragraph('_______________________________________________________________________________________')
doc.add_paragraph('_______________________________________________________________________________________')

doc.add_paragraph()
doc.add_paragraph().add_run('B4. Si vous gérez des copropriétés : comment réalisez-vous les appels de fonds ?').bold = True
doc.add_paragraph('□ Pas de syndic  □ Excel  □ Logiciel  □ Autre : ________________')

doc.add_paragraph()
doc.add_paragraph().add_run('B5. Comment gardez-vous une trace des visites et des contacts ?').bold = True
doc.add_paragraph('□ Papier/Cahier  □ Excel  □ WhatsApp  □ CRM  □ Autre : ________________')

doc.add_paragraph()
doc.add_paragraph().add_run('B6. Faites-vous des études de marché pour fixer vos prix ?').bold = True
doc.add_paragraph('□ Non, au jugé  □ Oui, via annonces en ligne  □ Oui, via réseau  □ Oui, via outil : ________________')

# Section C
doc.add_page_break()
doc.add_heading('SECTION C — ATTENTES SOLUTION', level=1)

doc.add_paragraph().add_run('C1. Quelle fonctionnalité serait un "game changer" pour votre agence ?').bold = True
doc.add_paragraph('_______________________________________________________________________________________')
doc.add_paragraph('_______________________________________________________________________________________')

doc.add_paragraph()
doc.add_paragraph().add_run('C2. Seriez-vous prêt à payer une solution de gestion ?').bold = True
doc.add_paragraph('□ Oui, budget mensuel : □ < 50 000  □ 50 000-100 000  □ 100 000-200 000  □ > 200 000 FCFA')
doc.add_paragraph('□ Non, je préfère une solution gratuite')

doc.add_paragraph()
doc.add_paragraph().add_run('C3. Seriez-vous prêt à tester gratuitement pendant 3 mois ?').bold = True
doc.add_paragraph('□ Oui, avec plaisir  □ Peut-être  □ Non')

doc.add_paragraph()
doc.add_paragraph().add_run('C4. Quels indicateurs suivriez-vous absolument sur un tableau de bord ?').bold = True
indicateurs = [
    'Taux d\'occupation',
    'Montant des impayés',
    'Rentabilité par bien',
    'Délai moyen de location',
    'Délai moyen de vente',
    'Évolution des prix/m²',
    'Activité des agents (visites, signatures)'
]
for ind in indicateurs:
    doc.add_paragraph(f'□ {ind}', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph().add_run('C5. Utilisez-vous des solutions de paiement mobile pour les loyers ?').bold = True
doc.add_paragraph('□ Non, uniquement espèces/virement')
doc.add_paragraph('□ Oui : □ Djamo  □ Orange Money  □ MTN Mobile Money  □ Wave  □ Autre : ________')

doc.add_paragraph()
doc.add_paragraph().add_run('C6. Quel serait le meilleur canal pour vous prévenir d\'une quittance impayée ?').bold = True
doc.add_paragraph('□ Email  □ SMS  □ WhatsApp  □ Notification dans l\'application  □ Appel téléphonique')

# Section D
doc.add_page_break()
doc.add_heading('SECTION D — OPEN FEEDBACK', level=1)

doc.add_paragraph().add_run('D1. Quel est le processus le plus chronophage dans votre agence actuellement ?').bold = True
doc.add_paragraph('_______________________________________________________________________________________')
doc.add_paragraph('_______________________________________________________________________________________')
doc.add_paragraph('_______________________________________________________________________________________')

doc.add_paragraph()
doc.add_paragraph().add_run('D2. Si vous pouviez supprimer une tâche administrative de votre quotidien, laquelle serait-ce ?').bold = True
doc.add_paragraph('_______________________________________________________________________________________')
doc.add_paragraph('_______________________________________________________________________________________')

doc.add_paragraph()
doc.add_paragraph().add_run('D3. Avez-vous des commentaires ou suggestions pour améliorer la gestion immobilière par un outil digital ?').bold = True
doc.add_paragraph('_______________________________________________________________________________________')
doc.add_paragraph('_______________________________________________________________________________________')
doc.add_paragraph('_______________________________________________________________________________________')

# Coordonnées
doc.add_page_break()
doc.add_heading('COORDONNÉES — ENTRETIEN SUIVI (optionnel)', level=1)
doc.add_paragraph('Si vous acceptez d\'être recontacté(e) pour un entretien approfondi de 30 min :')
doc.add_paragraph()
doc.add_paragraph('Nom du contact : _____________________________________________')
doc.add_paragraph('Téléphone : _____________________________________________')
doc.add_paragraph('Email : _____________________________________________')
doc.add_paragraph('Meilleur créneau : □ Matin  □ Midi  □ Après-midi')
doc.add_paragraph()
doc.add_paragraph('□ J\'accepte de recevoir les résultats agrégés de cette étude')

# Footer
doc.add_paragraph()
doc.add_paragraph('— Merci pour votre temps et votre contribution ! —').alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('/home/hdl/src/gestion_agence_immo/docs/Questionnaire_Agence.docx')
print("Questionnaire généré avec succès.")
